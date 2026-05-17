from __future__ import annotations

import logging
import uuid
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from sqlalchemy.orm import Session

from app.config import REPORT_DIR
from app.models.detection import DetectionRecord, Violation
from app.services import detection_service

logger = logging.getLogger(__name__)

VIOLATION_LABELS: dict[str, str] = {
    'warning_no_hardhat': '未戴安全帽',
    'warning_no_mask': '未戴口罩',
    'warning_no_safety_vest': '未穿反光背心',
    'warning_close_to_machinery': '人员靠近机械',
    'warning_close_to_vehicle': '人员靠近车辆',
    'warning_people_in_controlled_area': '人员进入管控区',
}


def generate_report(
    db: Session,
    start_date: datetime | None = None,
    end_date: datetime | None = None,
) -> str:
    """Generate a DOCX safety report.

    Args:
        db: Database session.
        start_date: Optional start date filter.
        end_date: Optional end date filter.

    Returns:
        Path to the generated report file.
    """
    if start_date and end_date:
        records = detection_service.get_records_by_date_range(
            db, start_date, end_date,
        )
    else:
        records = db.query(DetectionRecord).all()

    all_violations: list[Violation] = []
    for record in records:
        violations = detection_service.get_violations(db, record.id)
        all_violations.extend(violations)

    doc = Document()
    _add_title(doc, start_date, end_date)
    _add_summary(doc, records, all_violations)
    _add_violation_stats(doc, all_violations)
    _add_violation_screenshots(doc, all_violations)

    filename = f"report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    file_path = REPORT_DIR / filename
    doc.save(str(file_path))

    logger.info(f"Report generated: {file_path}")
    return str(file_path)


def _add_title(
    doc: Document,
    start_date: datetime | None,
    end_date: datetime | None,
) -> None:
    title = doc.add_heading('建筑施工现场安全隐患检测报告', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    date_range = '全部'
    if start_date and end_date:
        date_range = f"{start_date.strftime('%Y-%m-%d')} 至 {end_date.strftime('%Y-%m-%d')}"

    doc.add_paragraph(f"报告生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    doc.add_paragraph(f"数据范围: {date_range}")
    doc.add_paragraph()


def _add_summary(
    doc: Document,
    records: list[DetectionRecord],
    violations: list[Violation],
) -> None:
    doc.add_heading('一、检测概况', level=1)

    total_objects = sum(r.total_objects for r in records)
    total_violations = len(violations)

    table = doc.add_table(rows=5, cols=2, style='Light Shading Accent 1')
    table.style.font.size = Pt(11)

    summary_data = [
        ('检测记录总数', str(len(records))),
        ('检测目标总数', str(total_objects)),
        ('违规记录总数', str(total_violations)),
        ('视频总时长', f"{sum(r.duration for r in records):.1f} 秒"),
        ('报告生成时间', datetime.now().strftime('%Y-%m-%d %H:%M:%S')),
    ]

    for i, (label, value) in enumerate(summary_data):
        cell_label = table.cell(i, 0)
        cell_value = table.cell(i, 1)
        cell_label.text = label
        cell_value.text = value
        for paragraph in cell_label.paragraphs:
            paragraph.runs[0].bold = True


def _add_violation_stats(
    doc: Document,
    violations: list[Violation],
) -> None:
    doc.add_heading('二、违规类型统计', level=1)

    type_count: dict[str, int] = {}
    for v in violations:
        label = VIOLATION_LABELS.get(v.violation_type, v.violation_type)
        type_count[label] = type_count.get(label, 0) + 1

    if not type_count:
        doc.add_paragraph('暂无违规记录。')
        return

    table = doc.add_table(rows=1, cols=3, style='Light Shading Accent 1')
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '违规类型'
    hdr_cells[1].text = '次数'
    hdr_cells[2].text = '占比'

    for cell in hdr_cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True

    total = sum(type_count.values())
    for vtype, count in sorted(type_count.items(), key=lambda x: -x[1]):
        row = table.add_row()
        row.cells[0].text = vtype
        row.cells[1].text = str(count)
        row.cells[2].text = f"{count / total * 100:.1f}%"


def _add_violation_screenshots(
    doc: Document,
    violations: list[Violation],
) -> None:
    doc.add_heading('三、违规截图', level=1)

    screenshots = [v for v in violations if v.screenshot_path]

    if not screenshots:
        doc.add_paragraph('暂无违规截图。')
        return

    doc.add_paragraph(f"共 {len(screenshots)} 张违规截图：")

    for v in screenshots[:20]:  # Limit to 20 screenshots
        label = VIOLATION_LABELS.get(v.violation_type, v.violation_type)
        doc.add_paragraph()
        p = doc.add_paragraph()
        run = p.add_run(f"{label} - 帧 {v.frame_number} ({v.timestamp:.1f}s)")
        run.bold = True

        screenshot_full_path = Path(v.screenshot_path.lstrip('/'))
        if screenshot_full_path.exists():
            try:
                doc.add_picture(str(screenshot_full_path), width=Inches(4))
            except Exception:
                doc.add_paragraph(f"[图片加载失败: {screenshot_full_path}]")
